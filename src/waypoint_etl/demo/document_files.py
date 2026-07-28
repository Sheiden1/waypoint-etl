"""Documentos sintéticos de demonstração (TXT, DOCX e PDF digital).

Reproduzem o formato de fichas cadastrais e relatórios exportados por sistemas
antigos: rótulo e valor na mesma linha, para exercitar a extração por Regex.

Os dados vêm do mesmo gerador determinístico dos arquivos tabulares, portanto
são inteiramente sintéticos (seção 22).
"""

from __future__ import annotations

from pathlib import Path

import pymupdf
from docx import Document
from PIL import Image, ImageDraw

from .synthetic import FIXTURE_TIMESTAMP, generate_customer_rows

# Dimensões de uma folha A4 digitalizada a 150 DPI: resolução suficiente para
# o OCR sem gerar um fixture pesado.
SCAN_WIDTH = 1240
SCAN_HEIGHT = 1754
SCAN_MARGIN = 90
SCAN_LINE_HEIGHT = 42

# Campos exibidos na ficha, na ordem em que aparecem no documento.
FORM_FIELDS = (
    ("Código", "Código"),
    ("Nome", "Nome Cliente"),
    ("CPF/CNPJ", "CPF_CNPJ"),
    ("E-mail", "Correio Eletrônico"),
    ("Telefone", "Fone Principal"),
    ("CEP", "CEP"),
    ("Cidade", "Cidade"),
    ("UF", "UF"),
    ("Data de Cadastro", "Data Cadastro"),
)

FORM_TITLE = "FICHA CADASTRAL DE CLIENTE - ERP LEGADO"
TXT_TITLE = "RELATORIO DE CLIENTES - ERP LEGADO"
TXT_SEPARATOR = "-" * 62

_DEFAULT_FORM_COUNT = 3


def _form_lines(row: dict[str, str]) -> list[str]:
    """Monta as linhas ``Rótulo: valor`` de uma ficha."""
    return [f"{label}: {row[column]}".rstrip() for label, column in FORM_FIELDS]


def write_customers_txt(
    path: Path, *, seed: int = 20240101, count: int = 10
) -> Path:
    """Escreve um relatório de clientes em texto puro."""
    path.parent.mkdir(parents=True, exist_ok=True)
    rows = generate_customer_rows(seed=seed, count=count)[:count]

    lines = [TXT_TITLE, TXT_SEPARATOR]
    for row in rows:
        lines.extend(_form_lines(row))
        lines.append(TXT_SEPARATOR)

    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def write_customer_form_docx(
    path: Path, *, seed: int = 20240101, count: int = _DEFAULT_FORM_COUNT
) -> Path:
    """Escreve fichas cadastrais em Word, misturando parágrafos e tabela."""
    path.parent.mkdir(parents=True, exist_ok=True)
    rows = generate_customer_rows(seed=seed, count=count)[:count]

    document = Document()
    document.add_heading(FORM_TITLE, level=1)

    for row in rows:
        document.add_paragraph(f"Cliente {row['Código']}")
        # Uma tabela rótulo/valor, como nas fichas digitalizadas de verdade.
        table = document.add_table(rows=0, cols=2)
        for label, column in FORM_FIELDS:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = row[column]
        document.add_paragraph("")

    document.core_properties.created = FIXTURE_TIMESTAMP
    document.core_properties.modified = FIXTURE_TIMESTAMP
    document.save(str(path))
    return path


def write_customer_form_pdf(
    path: Path, *, seed: int = 20240101, count: int = _DEFAULT_FORM_COUNT
) -> Path:
    """Escreve fichas cadastrais em PDF com camada de texto (uma por página)."""
    path.parent.mkdir(parents=True, exist_ok=True)
    rows = generate_customer_rows(seed=seed, count=count)[:count]

    document = pymupdf.open()
    try:
        for row in rows:
            page = document.new_page()
            body = "\n".join([FORM_TITLE, "", *_form_lines(row)])
            # Margem de 72pt (uma polegada) em relação ao tamanho A4 padrão.
            box = pymupdf.Rect(72, 72, page.rect.width - 72, page.rect.height - 72)
            page.insert_textbox(box, body, fontsize=11, fontname="helv")
        stamp = FIXTURE_TIMESTAMP.strftime("D:%Y%m%d%H%M%SZ")
        document.set_metadata({"creationDate": stamp, "modDate": stamp})
        document.save(path)
    finally:
        document.close()
    return path


def render_form_image(
    *, seed: int = 20240101, index: int = 0
) -> Image.Image:
    """Desenha uma ficha cadastral como se fosse uma página digitalizada.

    Fonte bitmap padrão do Pillow, ampliada: não depende de nenhuma fonte
    instalada no sistema, o que manteria o fixture reprodutível em qualquer
    máquina.
    """
    row = generate_customer_rows(seed=seed, count=index + 1)[index]
    image = Image.new("L", (SCAN_WIDTH, SCAN_HEIGHT), color=255)
    draw = ImageDraw.Draw(image)

    y = SCAN_MARGIN
    draw.text((SCAN_MARGIN, y), FORM_TITLE, fill=0, font_size=34)
    y += SCAN_LINE_HEIGHT * 2

    for line in _form_lines(row):
        draw.text((SCAN_MARGIN, y), line, fill=0, font_size=30)
        y += SCAN_LINE_HEIGHT

    return image


def write_scanned_form_image(
    path: Path, *, seed: int = 20240101, index: int = 0
) -> Path:
    """Escreve uma ficha cadastral como imagem ("documento escaneado")."""
    path.parent.mkdir(parents=True, exist_ok=True)
    render_form_image(seed=seed, index=index).save(path)
    return path


def write_scanned_form_pdf(
    path: Path, *, seed: int = 20240101, count: int = 2
) -> Path:
    """Escreve um PDF **sem camada de texto**, só com a imagem das fichas.

    É o fixture que exercita o fallback para OCR: a extração nativa devolve
    zero caracteres, exatamente como um PDF escaneado de verdade.
    """
    path.parent.mkdir(parents=True, exist_ok=True)

    document = pymupdf.open()
    try:
        for index in range(count):
            image = render_form_image(seed=seed, index=index)
            page = document.new_page(width=595, height=842)
            page.insert_image(page.rect, pixmap=_to_pixmap(image))
        stamp = FIXTURE_TIMESTAMP.strftime("D:%Y%m%d%H%M%SZ")
        document.set_metadata({"creationDate": stamp, "modDate": stamp})
        document.save(path)
    finally:
        document.close()
    return path


def _to_pixmap(image: Image.Image) -> pymupdf.Pixmap:
    """Converte uma imagem do Pillow em um Pixmap do PyMuPDF."""
    rgb = image.convert("RGB")
    return pymupdf.Pixmap(
        pymupdf.csRGB, rgb.width, rgb.height, rgb.tobytes(), False
    )


__all__ = [
    "FORM_FIELDS",
    "FORM_TITLE",
    "TXT_TITLE",
    "render_form_image",
    "write_customer_form_docx",
    "write_customer_form_pdf",
    "write_customers_txt",
    "write_scanned_form_image",
    "write_scanned_form_pdf",
]
