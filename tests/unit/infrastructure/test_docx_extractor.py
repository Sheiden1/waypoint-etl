"""Testes do extrator de documentos Word."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from waypoint_etl.demo.document_files import FORM_TITLE, write_customer_form_docx
from waypoint_etl.domain.enums.source_format import SourceFormat
from waypoint_etl.domain.errors import ExtractionError, SourceNotFoundError
from waypoint_etl.infrastructure.extractors.docx_extractor import DocxExtractor


@pytest.fixture
def extractor() -> DocxExtractor:
    return DocxExtractor()


def test_supports_only_docx(extractor: DocxExtractor) -> None:
    assert extractor.supports(Path("a.docx"))
    assert extractor.supports(Path("A.DOCX"))
    assert not extractor.supports(Path("a.doc"))
    assert not extractor.supports(Path("a.txt"))


def test_reads_paragraphs(extractor: DocxExtractor, tmp_path: Path) -> None:
    source = tmp_path / "a.docx"
    document = Document()
    document.add_paragraph("Nome: Ana Silva")
    document.add_paragraph("CPF: 390.533.447-05")
    document.save(source)

    result = extractor.extract_text(source)

    assert result.source_format is SourceFormat.DOCX
    assert result.page_count == 1
    assert result.text == "Nome: Ana Silva\nCPF: 390.533.447-05"


def test_reads_tables_as_labelled_rows(
    extractor: DocxExtractor, tmp_path: Path
) -> None:
    source = tmp_path / "a.docx"
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "CPF"
    table.rows[0].cells[1].text = "390.533.447-05"
    document.save(source)

    result = extractor.extract_text(source)

    assert result.text == "CPF | 390.533.447-05"


def test_preserves_order_between_paragraphs_and_tables(
    extractor: DocxExtractor, tmp_path: Path
) -> None:
    """A ordem importa: o rótulo acima da tabela dá contexto ao valor."""
    source = tmp_path / "a.docx"
    document = Document()
    document.add_paragraph("Cliente ERP-1000")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "UF"
    table.rows[0].cells[1].text = "SP"
    document.add_paragraph("Cliente ERP-1001")
    document.save(source)

    result = extractor.extract_text(source)

    assert result.text.splitlines() == [
        "Cliente ERP-1000",
        "UF | SP",
        "Cliente ERP-1001",
    ]


def test_blank_paragraphs_are_dropped(
    extractor: DocxExtractor, tmp_path: Path
) -> None:
    source = tmp_path / "a.docx"
    document = Document()
    document.add_paragraph("Ana")
    document.add_paragraph("")
    document.add_paragraph("   ")
    document.add_paragraph("Bruno")
    document.save(source)

    result = extractor.extract_text(source)

    assert result.text.splitlines() == ["Ana", "Bruno"]


def test_empty_document_is_reported_but_not_fatal(
    extractor: DocxExtractor, tmp_path: Path
) -> None:
    source = tmp_path / "a.docx"
    Document().save(source)

    result = extractor.extract_text(source)

    assert result.is_empty
    assert any("não possui texto" in warning for warning in result.warnings)


def test_non_docx_content_raises_extraction_error(
    extractor: DocxExtractor, tmp_path: Path
) -> None:
    source = tmp_path / "a.docx"
    source.write_text("isto não é um documento Word", encoding="utf-8")

    with pytest.raises(ExtractionError, match="documento Word"):
        extractor.extract_text(source)


def test_missing_file_raises_source_not_found(
    extractor: DocxExtractor, tmp_path: Path
) -> None:
    with pytest.raises(SourceNotFoundError):
        extractor.extract_text(tmp_path / "inexistente.docx")


def test_demo_form_is_readable(extractor: DocxExtractor, tmp_path: Path) -> None:
    source = write_customer_form_docx(tmp_path / "ficha.docx")

    result = extractor.extract_text(source)

    assert FORM_TITLE in result.text
    assert "CPF/CNPJ |" in result.text
    assert not result.is_empty
